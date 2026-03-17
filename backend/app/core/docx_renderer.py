# backend/app/core/docx_renderer.py
"""
DOCX Renderer: Generates Microsoft Word documents from structured career data.
Follows ATS-safe best practices: single column, standard fonts, no tables/graphics.
"""

import io
from typing import Any

from docx import Document
from docx.shared import Pt

from app.core.theme_tokens import get_theme_tokens


def render_cover_letter_docx(
    content: str,
    candidate_name: str | None = None,
    theme_id: str = "minimal",
) -> bytes:
    """
    Renders a cover letter to DOCX format.
    """
    doc = Document()

    tokens = get_theme_tokens(theme_id)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = tokens["font"]["name"]
    font.size = Pt(tokens["font"]["size_pt"])

    # Header
    if candidate_name:
        h = doc.add_heading(candidate_name, level=0)
        h.alignment = tokens["alignment"]["heading"]

    # Body
    # Split content by newlines to handle paragraphs properly
    paragraphs = content.split("\n\n")
    for p_text in paragraphs:
        if p_text.strip():
            p = doc.add_paragraph(p_text.strip())
            p.alignment = tokens["alignment"]["body"]

    # Save to buffer
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()


def render_resume_docx(
    sections: dict[str, Any],
    candidate_name: str | None = None,
    theme_id: str = "minimal",
) -> bytes:
    """
    Renders a resume to DOCX format.
    Expects sections in a structured format (e.g. basics, work, education).
    """
    doc = Document()

    tokens = get_theme_tokens(theme_id)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = tokens["font"]["name"]
    font.size = Pt(tokens["font"]["size_pt"])

    # Header
    basics = sections.get("basics", {})
    name = candidate_name or basics.get("name", "Candidate")
    doc.add_heading(name, level=0)

    # Contact Info
    contact_parts = []
    if basics.get("email"):
        contact_parts.append(basics["email"])
    if basics.get("phone"):
        contact_parts.append(basics["phone"])
    if basics.get("location"):
        contact_parts.append(basics["location"])

    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.alignment = tokens["alignment"]["contact"]

    # Summary
    if sections.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(sections["summary"])

    # Work Experience
    if sections.get("work"):
        doc.add_heading("Work Experience", level=1)
        for job in sections["work"]:
            p = doc.add_paragraph()
            p.add_run(f"{job.get('role', 'Role')}").bold = True
            p.add_run(f" at {job.get('company', 'Company')}")

            # Dates
            dates = f"{job.get('startDate', '')} - {job.get('endDate', 'Present')}"
            p.add_run(f"\t{dates}").italic = True

            # Bullets
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if sections.get("education"):
        doc.add_heading("Education", level=1)
        for edu in sections["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('qualification', 'Degree')}").bold = True
            p.add_run(f", {edu.get('institution', 'Institution')}")
            if edu.get("endDate"):
                p.add_run(f"\t{edu.get('endDate')}").italic = True

    # Skills
    if sections.get("skills"):
        doc.add_heading("Skills", level=1)
        # Handle both list of strings and groups
        if isinstance(sections["skills"], list):
            doc.add_paragraph(", ".join(sections["skills"]))
        elif isinstance(sections["skills"], dict):
            for group, items in sections["skills"].items():
                p = doc.add_paragraph()
                p.add_run(f"{group}: ").bold = True
                p.add_run(", ".join(items))

    # Save to buffer
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()


def render_ksc_docx(
    responses: list[dict[str, Any]],
    job_title: str = "Job Application",
    theme_id: str = "minimal",
) -> bytes:
    """
    Renders Key Selection Criteria responses to DOCX.
    """
    doc = Document()

    tokens = get_theme_tokens(theme_id)

    style = doc.styles["Normal"]
    font = style.font
    font.name = tokens["font"]["name"]
    font.size = Pt(tokens["font"]["size_pt"])

    h = doc.add_heading(f"Selection Criteria: {job_title}", level=0)
    h.alignment = tokens["alignment"]["heading"]

    for item in responses:
        criterion = item.get("criterion", "Criterion")
        response = item.get("response", "")

        doc.add_heading(criterion, level=1)

        # Split by STAR if present or just paragraphs
        paragraphs = response.split("\n\n")
        for p_text in paragraphs:
            if p_text.strip():
                doc.add_paragraph(p_text.strip())

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()
