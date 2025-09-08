import io
import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional

import docx
import pdfplumber
from app.core.dependencies import get_current_user_with_state, get_user_document_from_firestore
from app.core.file_upload_decorators import require_valid_resume_upload, require_valid_document_upload
from app.core.limiter import authenticated_limiter
from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile, status
from google.api_core.exceptions import GoogleAPICallError
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel
from starlette.responses import StreamingResponse

# Temporarily disabled weasyprint import due to missing system dependencies
# try:
#     from weasyprint import CSS, HTML
#     WEASYPRINT_AVAILABLE = True
# except ImportError:
WEASYPRINT_AVAILABLE = False
CSS = None
HTML = None

# from app.genkit_flows.extract_resume_entities import
# extract_resume_entities  # Temporarily disabled for deployment

router = APIRouter()
logger = logging.getLogger(__name__)


class Document(BaseModel):
    id: str
    name: str
    type: str  # 'resume', 'cover_letter', etc.
    size: int
    created_at: str
    updated_at: str
    content_preview: Optional[str] = None
    status: str = "active"
    metadata: Optional[Dict[str, Any]] = None


# --- Template Configuration Loading ---
config_path = Path(__file__).parent.parent.parent.parent / "config" / "themes.json"
template_root_dir = Path(__file__).parent.parent.parent / "templates"

with open(config_path) as f:
    THEME_CONFIG = json.load(f)

# Dynamically create the Theme literal type from the config file keys
Theme = Literal[tuple(THEME_CONFIG.keys())]
env = Environment(loader=FileSystemLoader(str(template_root_dir)))


def parse_pdf(file_path: str) -> str:
    with pdfplumber.open(file_path) as pdf:
        return "".join(page.extract_text() for page in pdf.pages if page.extract_text())


def parse_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\\n".join(para.text for para in doc.paragraphs)


async def process_and_upload_file(file: UploadFile, uid: str, doc_type: str):
    # This is a placeholder for the actual implementation
    pass


@router.post("/upload")
@authenticated_limiter.limit("20/minute")  # Allow more uploads but still rate limit
@require_valid_document_upload(max_files=5)  # Allow up to 5 documents, validate each
async def upload_and_parse_files(
    request: Request,
    files: List[UploadFile] = File(...),
    user: dict = Depends(get_current_user_with_state),
    doc_type: str = "resume",
):
    """Upload and parse multiple documents with validation.
    
    This endpoint now uses the file upload validation decorator to:
    - Validate file types and extensions
    - Check file sizes
    - Ensure filenames are safe
    - Limit the number of files
    """
    # Files are already validated by the decorator
    uploaded_documents = []
    
    for file in files:
        try:
            # Process each validated file
            result = await process_and_upload_file(file, user["uid"], doc_type)
            uploaded_documents.append({
                "filename": file.filename,
                "status": "success",
                "result": result
            })
        except Exception as e:
            logger.error(f"Failed to process file {file.filename}: {str(e)}")
            uploaded_documents.append({
                "filename": file.filename,
                "status": "error",
                "error": str(e)
            })
    
    return {
        "message": f"Processed {len(files)} files",
        "results": uploaded_documents,
        "total_files": len(files),
        "successful": len([r for r in uploaded_documents if r["status"] == "success"]),
        "failed": len([r for r in uploaded_documents if r["status"] == "error"])
    }


@router.get("")
@authenticated_limiter.limit("30/minute")
async def list_documents(
    request: Request,
    user: dict = Depends(get_current_user_with_state),
    doc_type: Optional[str] = None,
) -> List[Document]:
    """List all documents for the current user"""
    user_id = user.get("uid") or user.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User ID not found")

    # Mock data for now - in production this would come from database/Firestore
    mock_documents = [
        Document(
            id="doc1",
            name="Software Engineer Resume.pdf",
            type="resume",
            size=245760,  # ~240KB
            created_at="2024-01-15T10:30:00Z",
            updated_at="2024-01-15T10:30:00Z",
            content_preview="Experienced Software Engineer with 5+ years...",
            status="active",
            metadata={"pages": 2, "format": "pdf"},
        ),
        Document(
            id="doc2",
            name="Cover Letter - Tech Company.docx",
            type="cover_letter",
            size=51200,  # ~50KB
            created_at="2024-01-20T14:15:00Z",
            updated_at="2024-01-20T14:15:00Z",
            content_preview="Dear Hiring Manager, I am excited to apply...",
            status="active",
            metadata={"pages": 1, "format": "docx"},
        ),
    ]

    # Filter by document type if specified
    if doc_type:
        mock_documents = [doc for doc in mock_documents if doc.type == doc_type]

    return mock_documents


@router.get("/{document_id}/download-pdf")
async def download_document_as_pdf(
    document_id: str,
    theme: Theme,
    document: dict = Depends(get_user_document_from_firestore),
):
    try:
        content = document.get("content", "")
        theme_details = THEME_CONFIG.get(theme)

        if not theme_details:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Theme '{theme}' not found.",
            )

        template_path = theme_details["template"]
        css_path = str(template_root_dir / theme_details["css"])
        base_url_path = str(template_root_dir / theme)

        if not WEASYPRINT_AVAILABLE:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="PDF generation temporarily unavailable - weasyprint dependencies missing",
            )

        template = env.get_template(template_path)
        html_content = template.render(content=content)
        stylesheet = CSS(filename=css_path)
        pdf_bytes = HTML(string=html_content, base_url=base_url_path).write_pdf(
            stylesheets=[stylesheet]
        )

        response = StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf")
        original_filename = document.get("originalFilename", "document").split(".")[0]
        response.headers[
            "Content-Disposition"
        ] = f"attachment; filename={original_filename}_{theme}.pdf"

        return response

    except (FileNotFoundError, GoogleAPICallError) as e:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=f"Error accessing template files or cloud storage: {e}",
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An error occurred while generating the PDF: {e}",
        )
