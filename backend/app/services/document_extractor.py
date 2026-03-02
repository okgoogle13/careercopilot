"""
document_extractor.py

Service for extracting text from PDF and Word documents.
Used when job descriptions or selection criteria are in attached documents.

Supports:
- PDF files (via pypdf + pdfminer fallback)
- Word documents (via python-docx if installed)
- Automatic detection of document links on job pages
"""

import logging
import re
<<<<<<< HEAD
from typing import Optional, List, Tuple
from urllib.parse import urljoin, urlparse
import requests
from io import BytesIO
=======
from io import BytesIO
from urllib.parse import urljoin

import requests
>>>>>>> restoration-KR-Rage-Figma-v2.0

logger = logging.getLogger(__name__)

# PDF extraction
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("pypdf not available for PDF extraction")

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False
    logger.warning("pdfminer not available for PDF extraction fallback")

# Word document extraction
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.info("python-docx not available for Word document extraction")


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file content.

    Args:
        file_content: PDF file as bytes

    Returns:
        Extracted text content
    """
    text = ""

    # Try pypdf first (faster)
    if PYPDF_AVAILABLE:
        try:
            pdf_file = BytesIO(file_content)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"

            if text.strip():
                logger.info(f"Extracted {len(text)} characters from PDF using pypdf")
                return text
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}, trying pdfminer")

    # Fallback to pdfminer (more robust)
    if PDFMINER_AVAILABLE:
        try:
            pdf_file = BytesIO(file_content)
            text = pdfminer_extract(pdf_file)
            logger.info(f"Extracted {len(text)} characters from PDF using pdfminer")
            return text
        except Exception as e:
            logger.error(f"pdfminer extraction failed: {e}")

    raise RuntimeError("PDF extraction failed with all available methods")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word document.

    Args:
        file_content: DOCX file as bytes

    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed, cannot extract Word documents")

    try:
        doc_file = BytesIO(file_content)
        doc = Document(doc_file)

        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        logger.info(f"Extracted {len(text)} characters from Word document")
        return text
    except Exception as e:
        logger.error(f"Word document extraction failed: {e}")
<<<<<<< HEAD
        raise RuntimeError(f"Failed to extract Word document: {str(e)}")


def detect_document_links(html_content: str, base_url: str) -> List[Tuple[str, str]]:
=======
        raise RuntimeError(f"Failed to extract Word document: {e!s}")


def detect_document_links(html_content: str, base_url: str) -> list[tuple[str, str]]:
>>>>>>> restoration-KR-Rage-Figma-v2.0
    """
    Detect PDF and Word document links in HTML content.

    Args:
        html_content: HTML page content
        base_url: Base URL for resolving relative links

    Returns:
        List of (url, type) tuples where type is 'pdf' or 'docx'
    """
    documents = []

    # Patterns for document links
    pdf_pattern = r'href=["\']([^"\']*\.pdf[^"\']*)["\']'
    docx_pattern = r'href=["\']([^"\']*\.docx?[^"\']*)["\']'

    # Find PDF links
    for match in re.finditer(pdf_pattern, html_content, re.IGNORECASE):
        url = match.group(1)
        absolute_url = urljoin(base_url, url)
<<<<<<< HEAD
        documents.append((absolute_url, 'pdf'))
=======
        documents.append((absolute_url, "pdf"))
>>>>>>> restoration-KR-Rage-Figma-v2.0

    # Find Word document links
    for match in re.finditer(docx_pattern, html_content, re.IGNORECASE):
        url = match.group(1)
        absolute_url = urljoin(base_url, url)
<<<<<<< HEAD
        documents.append((absolute_url, 'docx'))
=======
        documents.append((absolute_url, "docx"))
>>>>>>> restoration-KR-Rage-Figma-v2.0

    logger.info(f"Detected {len(documents)} document links: {documents}")
    return documents


def download_document(url: str, timeout: int = 30) -> bytes:
    """
    Download document from URL.

    Args:
        url: Document URL
        timeout: Request timeout in seconds

    Returns:
        Document content as bytes
    """
    try:
        response = requests.get(url, timeout=timeout, headers={
<<<<<<< HEAD
            'User-Agent': 'Mozilla/5.0 (compatible; JobAnalyzer/1.0)'
=======
            "User-Agent": "Mozilla/5.0 (compatible; JobAnalyzer/1.0)"
>>>>>>> restoration-KR-Rage-Figma-v2.0
        })
        response.raise_for_status()
        return response.content
    except Exception as e:
        logger.error(f"Failed to download document from {url}: {e}")
<<<<<<< HEAD
        raise RuntimeError(f"Document download failed: {str(e)}")
=======
        raise RuntimeError(f"Document download failed: {e!s}")
>>>>>>> restoration-KR-Rage-Figma-v2.0


def extract_documents_from_page(html_content: str, base_url: str) -> str:
    """
    Detect and extract text from all documents linked on a page.

    Args:
        html_content: HTML page content
        base_url: Base URL for the page

    Returns:
        Combined text from all extracted documents
    """
    documents = detect_document_links(html_content, base_url)

    if not documents:
        logger.info("No document attachments found on page")
        return ""

    combined_text = ""

    for doc_url, doc_type in documents:
        try:
            logger.info(f"Downloading {doc_type} from {doc_url}")
            content = download_document(doc_url)

<<<<<<< HEAD
            if doc_type == 'pdf':
                text = extract_text_from_pdf(content)
            elif doc_type == 'docx':
=======
            if doc_type == "pdf":
                text = extract_text_from_pdf(content)
            elif doc_type == "docx":
>>>>>>> restoration-KR-Rage-Figma-v2.0
                text = extract_text_from_docx(content)
            else:
                logger.warning(f"Unknown document type: {doc_type}")
                continue

            combined_text += f"\n\n--- Document: {doc_url} ---\n{text}\n"

        except Exception as e:
            logger.warning(f"Failed to extract {doc_type} from {doc_url}: {e}")
            # Continue with other documents

    return combined_text
